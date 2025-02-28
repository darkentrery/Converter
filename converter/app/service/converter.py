import base64
import io
import os
import subprocess
import tempfile
from pathlib import Path

import pandas as pd
import pdfkit
from PIL import Image
from docx import Document
from docx.shared import Inches
from lxml import etree
from pdf2docx import Converter

from app import entity
from app.config import config
from app.logger import logger
from app.repository.sqlalchemy import SAUnitOfWork


class ConverterService:
    def __init__(self, uow: SAUnitOfWork):
        self.uow = uow

    async def from_jpg_to_pdf(self, orientation: str, images: list[bytes]) -> io.BytesIO:
        images = [self.resize_to_a4(Image.open(io.BytesIO(img)).convert("RGB"), orientation) for img in images]
        pdf_bytes = io.BytesIO()
        images[0].save(pdf_bytes, format="PDF", save_all=True, append_images=images[1:])
        pdf_bytes.seek(0)
        return pdf_bytes

    def resize_to_a4(self, image: Image.Image, orientation: str) -> Image.Image:
        """Масштабирует изображение под размер A4, сохраняя пропорции."""
        orientation = orientation if orientation in [entity.Orientation.PORTRAIT.value,
                                                     entity.Orientation.LANDSCAPE.value] else self.get_orientation_depends_size(image)
        size = config.A4_LANDSCAPE if orientation == entity.Orientation.LANDSCAPE.value else config.A4_PORTRAIT
        image.thumbnail(size, Image.Resampling.LANCZOS)
        new_img = Image.new("RGB", size, "white")
        x_offset = (size[0] - image.width) // 2
        y_offset = (size[1] - image.height) // 2
        new_img.paste(image, (x_offset, y_offset))
        return new_img

    def get_orientation_depends_size(self, image: Image) -> str:
        if image.width > image.height:
            return entity.Orientation.LANDSCAPE.value
        else:
            return entity.Orientation.PORTRAIT.value

    def from_word_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        return self._convert_by_libreoffice(files, "docx", "pdf")

    def from_powerpoint_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        return self._convert_by_libreoffice(files, "pptx", "pdf")

    def from_txt_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        return self._convert_by_libreoffice(files, "txt", "pdf")

    def from_pdf_to_word(self, files: list[bytes]) -> io.BytesIO:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            input_path = Path(temp_file.name)
            output_path = input_path.with_suffix(f".docx")

        cv = Converter(stream=files[0])
        cv.convert(str(output_path), start=0, end=None)
        cv.close()

        with open(output_path, "rb") as pdf_file:
            word_bytes = pdf_file.read()

        input_path.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)

        return io.BytesIO(word_bytes)

    def _convert_by_libreoffice(self, files: list[bytes], from_suffix: str, to_suffix: str) -> io.BytesIO:
        """Конвертирует документ в PDF используя LibreOffice (без сохранения на диск)"""

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{from_suffix}") as temp_file:
            logger.info(f"{temp_file.name=}")
            temp_file.write(files[0])
            temp_file.flush()
            temp_path = Path(temp_file.name)
            output_pdf_path = temp_path.with_suffix(f".{to_suffix}")

        # Запускаем LibreOffice для конвертации
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", to_suffix,
            str(temp_path), "--outdir", str(temp_path.parent)
        ], check=True)

        # Читаем PDF в память
        with open(output_pdf_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()

        # Удаляем временные файлы
        temp_path.unlink(missing_ok=True)
        output_pdf_path.unlink(missing_ok=True)

        return io.BytesIO(pdf_bytes)

    def from_excel_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует Excel (bytes) в PDF (bytes)"""
        file_stream = io.BytesIO(files[0])
        df = pd.read_excel(file_stream)
        html_content = df.to_html(index=False, border=1)
        return self._convert_by_pdfkit(html_content, "pdf")

    def from_csv_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует Excel (bytes) в PDF (bytes)"""
        file_stream = io.BytesIO(files[0])
        df = pd.read_csv(file_stream)
        html_content = df.to_html(index=False, border=1)
        return self._convert_by_pdfkit(html_content, "pdf")

    def from_html_to_pdf(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует Html (bytes) в PDF (bytes)"""
        html_content = files[0].decode('utf-8')
        return self._convert_by_pdfkit(html_content, "pdf")

    def _convert_by_pdfkit(self, content: str, to_suffix: str) -> io.BytesIO:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{to_suffix}") as temp_file:
            temp_path = temp_file.name

        options = {
            "page-size": "A4",
            "encoding": "UTF-8",
            "no-outline": None
        }

        pdfkit.from_string(content, temp_path, options=options)
        with open(temp_path, 'rb') as f:
            pdf_stream = io.BytesIO(f.read())

        os.remove(temp_path)

        pdf_stream.seek(0)
        return pdf_stream

    def from_jpg_to_word(self, files: list[bytes]) -> io.BytesIO:
        doc = Document()
        # doc.add_heading("Images to Word", level=1)  # Заголовок

        # Добавляем изображения в документ
        for idx, img_bytes in enumerate(files):
            # Читаем изображение из байтов
            img_stream = io.BytesIO(img_bytes)
            # doc.add_paragraph(f"Image {idx + 1}")  # Подпись к изображению
            doc.add_picture(img_stream, width=Inches(5))  # Добавляем изображение (ширина 5 дюймов)
            doc.add_paragraph("\n")  # Отступ

        # Сохраняем документ в памяти с помощью BytesIO
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)  # Сброс позиции потока, чтобы его можно было прочитать

        return doc_stream

    def from_csv_to_excel(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует Excel (bytes) в PDF (bytes)"""
        file_stream = io.BytesIO(files[0])
        df = pd.read_csv(file_stream)

        output_buffer = io.BytesIO()
        df.to_excel(output_buffer, index=False)
        output_buffer.seek(0)
        return output_buffer

    def from_excel_to_csv(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует Excel (bytes) в PDF (bytes)"""
        file_stream = io.BytesIO(files[0])
        df = pd.read_excel(file_stream)

        output_buffer = io.BytesIO()
        df.to_csv(output_buffer, index=False, encoding="utf-8")
        output_buffer.seek(0)
        return output_buffer

    def from_pdf_to_html(self, files: list[bytes]) -> io.BytesIO:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".pdf") as temp_file:
            temp_file.write(files[0])
            temp_file.flush()
            input_path = Path(temp_file.name)
            output_path = input_path.with_suffix(f".html")

        # Запускаем pdf2htmlEX
        subprocess.run(
            ["docker", "exec", "pdf2html", "pdf2htmlEX", f"{input_path}", f"{output_path.name}"],
            check=True, capture_output=True
        )
        with open(output_path, "rb") as html_file:
            html_content = io.BytesIO(html_file.read())
            html_content.seek(0)

        os.remove(input_path)
        os.remove(output_path)

        return html_content

    def from_word_to_fb2(self, files: list[bytes]) -> io.BytesIO:
        """Конвертирует .docx (переданный как байты) в .fb2"""
        doc = Document(io.BytesIO(files[0]))

        fb2 = etree.Element("FictionBook", xmlns="http://www.gribuser.ru/xml/fictionbook/2.0")
        description = etree.SubElement(fb2, "description")
        title_info = etree.SubElement(description, "title-info")

        book_title = etree.SubElement(title_info, "book-title")
        book_title.text = "Название книги"

        body = etree.SubElement(fb2, "body")
        body_content = etree.SubElement(body, "section")

        img_counter = 1  # Счетчик изображений
        images_dict = {}  # Словарь для хранения изображений

        def add_paragraph(text: str) -> None:
            """Добавляет параграф в FB2"""
            p = etree.SubElement(body_content, "p")
            p.text = text

        def add_image(image_bytes: bytes) -> None:
            """Добавляет картинку в тело FB2 и в <binary>"""
            nonlocal img_counter
            img_id = f"img_{img_counter}"
            img_counter += 1

            # Кодируем картинку в base64
            encoded_image = base64.b64encode(image_bytes).decode("utf-8")

            # Добавляем <image> в текст
            etree.SubElement(body_content, "image", href=f"#{img_id}")

            # Сохраняем картинку в словарь для <binary>
            images_dict[img_id] = encoded_image

        def process_element(element):
            """Обрабатывает элементы документа (текст и картинки)"""
            if element.tag.endswith("p"):  # Параграф
                text = element.text or ""
                add_paragraph(text.strip())

            elif element.tag.endswith("tbl"):  # Таблица
                for row in element.findall(".//{*}tr"):  # Строки таблицы
                    row_text = []
                    for cell in row.findall(".//{*}tc"):  # Ячейки
                        paragraphs = cell.findall(".//{*}p")  # Ищем только параграфы в ячейке
                        cell_text = " ".join(
                            p.text.strip() for p in paragraphs if p.text)  # Склеиваем текст в одной ячейке
                        row_text.append(cell_text)

                        # Добавляем строку таблицы как один параграф
                    if row_text:
                        add_paragraph(" | ".join(row_text))

            # Извлекаем картинки прямо из параграфов и таблиц
            for shape in element.findall(".//{*}graphic"):
                try:
                    pic = shape.find(".//{*}pic")
                    if pic is not None:
                        blip_fill = pic.find(".//{*}blipFill")
                        if blip_fill is not None:
                            blip = blip_fill.find(".//{*}blip")
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")

                            # Получаем изображение
                            image_part = doc.part.related_parts[embed_id]
                            add_image(image_part.blob)
                except Exception as e:
                    logger.error(f"Ошибка извлечения изображения: {e}")

        # Обрабатываем основной контент
        for element in doc.element.body:
            process_element(element)

        # Добавляем изображения в <binary>
        for img_id, img_data in images_dict.items():
            binary = etree.SubElement(fb2, "binary", id=img_id, content_type="image/jpeg")
            binary.text = img_data

        content = etree.tostring(fb2, pretty_print=True, encoding="utf-8", xml_declaration=True)
        return io.BytesIO(content)
